"""Text extraction service — used by both the upload endpoint and the CLI script.

Returns (text, used_ocr) so the caller can flag OCR-derived documents.
"""
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

import structlog

from app.services import ocr as ocr_service

log = structlog.get_logger()

# `.doc` is supported via LibreOffice headless conversion → `.docx`, then
# the normal docx path. Requires libreoffice-core + libreoffice-writer to
# be installed on the host (see render.yaml buildCommand). If soffice is
# missing at runtime, .doc uploads return a clear ExtractionResult with a
# note instead of crashing.
SUPPORTED_EXTENSIONS = {".txt", ".md", ".doc", ".docx", ".pdf"}


@dataclass
class ExtractionResult:
    text: str
    used_ocr: bool
    extractor: str  # which path was taken: "txt" | "docx" | "pdfplumber" | "azure_ocr"
    note: str | None = None  # human-readable diagnostic, e.g. "scanned PDF — no OCR configured"
    pages: int | None = None  # source page count when known (PDFs)
    failed_ocr_batches: list[int] | None = None  # set when OCR partially failed
    partial: bool = False  # True when extractor returned incomplete content
    quality: str = "ok"  # ok | low_density | partial | reversed_hebrew | empty


# Common Hebrew function words. If we see more reversed forms than forward
# forms in extracted text, the PDF text layer is RTL-visual-ordered garbage.
_HE_FORWARD = ("של", "את", "על", "לא", "כי", "אם", "זה", "אל", "כל", "מן", "או")
_HE_REVERSED = tuple(w[::-1] for w in _HE_FORWARD)


def _looks_reversed_hebrew(text: str) -> bool:
    """Heuristic: count reversed vs forward Hebrew function words.

    pdfplumber on some scanned Hebrew PDFs returns RTL text in visual order —
    each line is character-by-character LTR, so words come out backwards.
    Density check passes (chars are present), but the content is unusable.
    """
    if not text or len(text) < 200:
        return False
    fwd = sum(text.count(w) for w in _HE_FORWARD)
    rev = sum(text.count(w) for w in _HE_REVERSED)
    return rev > fwd + 3  # require margin to avoid false positives on short docs


def extract_text(path: Path, prefer_ocr: bool = False) -> ExtractionResult:
    """Extract text from a file. Falls back to Azure OCR if a PDF has no native text.

    prefer_ocr: for PDFs, skip pdfplumber and go straight to Azure OCR. Useful
    for scanned Hebrew PDFs where pdfplumber returns reversed RTL text that
    looks valid (passes density check) but is unusable for embedding/search.
    """
    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return ExtractionResult(text=path.read_text(encoding="utf-8"), used_ocr=False, extractor="txt")

    if suffix == ".docx":
        return ExtractionResult(text=_extract_docx(path), used_ocr=False, extractor="docx")

    if suffix == ".doc":
        # Legacy binary Word format. python-docx can't read it; we shell out
        # to LibreOffice headless to convert to .docx, then reuse the docx
        # path. The alternative Python-only libs (antiword, olefile) handle
        # Hebrew poorly; LibreOffice's docx exporter is the same one Office
        # itself uses via the "Save As" round-trip.
        converted = _convert_doc_to_docx(path)
        if converted is None:
            return ExtractionResult(
                text="",
                used_ocr=False,
                extractor="doc_unconvertible",
                note=(
                    "קובץ .doc — נדרשת LibreOffice להמרה ל-.docx, אך היא לא "
                    "מותקנת בשרת. שמור את הקובץ מחדש כ-.docx ונסה שוב."
                ),
            )
        try:
            return ExtractionResult(
                text=_extract_docx(converted),
                used_ocr=False,
                extractor="doc_via_libreoffice",
            )
        finally:
            try:
                converted.unlink()
                # soffice writes into a temp dir; clean the whole dir too.
                converted.parent.rmdir()
            except OSError:
                pass

    if suffix == ".pdf":
        page_count = _pdf_page_count(path)
        if not prefer_ocr:
            native = _extract_pdf_native(path)
            if native.strip():
                # Detect RTL-visual-order garbage — if so, escalate to OCR
                # instead of accepting reversed Hebrew that would silently
                # poison embeddings + search.
                if _looks_reversed_hebrew(native) and ocr_service.is_configured():
                    log.warning("extraction.pdfplumber_reversed_hebrew_falling_back_to_ocr",
                                path=str(path))
                else:
                    return ExtractionResult(
                        text=native, used_ocr=False, extractor="pdfplumber", pages=page_count,
                    )

        if not ocr_service.is_configured():
            return ExtractionResult(
                text="",
                used_ocr=False,
                extractor="pdfplumber",
                pages=page_count,
                note="PDF has no extractable text (likely scanned). Azure OCR not configured.",
            )

        try:
            ocr_text = ocr_service.ocr_pdf(path)
        except ocr_service.PartialOcrError as e:
            return ExtractionResult(
                text=e.partial_text,
                used_ocr=True,
                extractor="azure_ocr",
                pages=page_count,
                failed_ocr_batches=e.failed_batches,
                partial=True,
                note=(
                    f"Azure OCR failed on {len(e.failed_batches)}/{e.total_batches} batches "
                    f"({e.failed_batches}); partial text returned."
                ),
            )
        return ExtractionResult(
            text=ocr_text,
            used_ocr=True,
            extractor="azure_ocr",
            pages=page_count,
            note=None if ocr_text.strip() else "Azure OCR returned no text",
        )

    return ExtractionResult(
        text="",
        used_ocr=False,
        extractor="unsupported",
        note=f"Unsupported file type: {suffix}",
    )


def _convert_doc_to_docx(path: Path) -> Path | None:
    """Convert a legacy .doc file to .docx via LibreOffice headless.

    Returns the path to the produced .docx (inside a fresh temp dir the
    caller must delete), or None if soffice isn't installed or the
    conversion failed. Bounded 60s timeout — a stuck soffice would
    otherwise wedge the ingest request until Render's healthcheck kills
    the worker.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        log.warning("extraction.doc_soffice_missing", path=str(path))
        return None
    outdir = Path(tempfile.mkdtemp(prefix="doc2docx_"))
    try:
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(outdir),
                str(path),
            ],
            capture_output=True,
            text=True,
            timeout=60,
            check=False,
        )
    except subprocess.TimeoutExpired:
        log.warning("extraction.doc_conversion_timeout", path=str(path))
        try:
            outdir.rmdir()
        except OSError:
            pass
        return None
    if result.returncode != 0:
        log.warning(
            "extraction.doc_conversion_failed",
            path=str(path),
            stderr=(result.stderr or "")[:300],
        )
        try:
            outdir.rmdir()
        except OSError:
            pass
        return None
    # LibreOffice writes <stem>.docx into outdir regardless of input name.
    produced = outdir / f"{path.stem}.docx"
    if not produced.exists():
        # Fallback: pick the first .docx it wrote (some builds mangle stems
        # containing non-ASCII characters).
        candidates = list(outdir.glob("*.docx"))
        if not candidates:
            return None
        produced = candidates[0]
    return produced


def _extract_docx(path: Path) -> str:
    from docx import Document  # type: ignore[import-not-found]

    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def _extract_pdf_native(path: Path) -> str:
    import pdfplumber  # type: ignore[import-not-found]

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
    return "\n\n".join(pages)


def _pdf_page_count(path: Path) -> int | None:
    try:
        import pymupdf  # type: ignore[import-not-found]

        with pymupdf.open(path) as doc:
            return doc.page_count
    except Exception:
        return None
