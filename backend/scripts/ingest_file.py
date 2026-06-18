"""Ingest a real document file (or folder) into the platform.

Supports: .txt, .md, .docx, .pdf
For PDFs: works for native PDFs (extractable text). Scanned PDFs need OCR
(not wired up yet).

Usage:
    cd backend
    .venv/bin/python -m scripts.ingest_file <path-to-file-or-folder> [options]

Options:
    --doc-type TYPE        bylaw | sub_bylaw | minutes | decision | other (default: bylaw)
    --tenant-id UUID       which tenant to ingest into (default: first tenant)
    --api URL              backend URL (default: http://localhost:8000)
    --dry-run              print extracted text without ingesting
"""
import argparse
import json
import sys
from pathlib import Path
from urllib import request

SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")

    if suffix == ".docx":
        from docx import Document  # type: ignore[import-not-found]

        doc = Document(path)
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract tables (common in protocols / bylaws)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts)

    if suffix == ".pdf":
        import pdfplumber  # type: ignore[import-not-found]

        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(text)
                else:
                    print(f"  ⚠ page {i} produced no text (likely scanned — OCR needed)", file=sys.stderr)
        return "\n\n".join(pages)

    raise ValueError(f"Unsupported file type: {suffix}")


def ingest_one(path: Path, *, doc_type: str, tenant_id: str | None, api: str, dry_run: bool) -> None:
    print(f"\n▶ {path.name}")
    try:
        text = extract_text(path)
    except Exception as e:
        print(f"  ✗ extraction failed: {e}", file=sys.stderr)
        return

    n_chars = len(text)
    n_words = len(text.split())
    print(f"  extracted {n_chars} chars / ~{n_words} words")

    if not text.strip():
        print("  ✗ no text content; skipping", file=sys.stderr)
        return

    if dry_run:
        print(f"  --- preview ---\n{text[:600]}{'…' if len(text) > 600 else ''}\n  ---")
        return

    body: dict[str, object] = {
        "filename": path.name,
        "text": text,
        "doc_type": doc_type,
    }
    if tenant_id:
        body["tenant_id"] = tenant_id

    data = json.dumps(body, ensure_ascii=False).encode("utf-8")
    req = request.Request(
        f"{api.rstrip('/')}/api/ingest",
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with request.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            print(f"  ✓ document_id={result['document_id']}  chunks_created={result['chunks_created']}")
    except Exception as e:
        print(f"  ✗ ingest failed: {e}", file=sys.stderr)


def iter_targets(root: Path) -> list[Path]:
    if root.is_file():
        return [root] if root.suffix.lower() in SUPPORTED_EXTENSIONS else []
    if root.is_dir():
        return sorted(
            p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
        )
    return []


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("path", type=Path, help="File or folder to ingest")
    parser.add_argument("--doc-type", default="bylaw", choices=["bylaw", "sub_bylaw", "minutes", "decision", "other"])
    parser.add_argument("--tenant-id", default=None)
    parser.add_argument("--api", default="http://localhost:8000")
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    if not args.path.exists():
        print(f"Path not found: {args.path}", file=sys.stderr)
        sys.exit(1)

    targets = iter_targets(args.path)
    if not targets:
        print(f"No supported files found at {args.path}", file=sys.stderr)
        print(f"Supported extensions: {', '.join(sorted(SUPPORTED_EXTENSIONS))}", file=sys.stderr)
        sys.exit(1)

    print(f"Found {len(targets)} file(s) to ingest:")
    for t in targets:
        print(f"  · {t}")

    for t in targets:
        ingest_one(
            t,
            doc_type=args.doc_type,
            tenant_id=args.tenant_id,
            api=args.api,
            dry_run=args.dry_run,
        )

    print("\nDone.")


if __name__ == "__main__":
    main()
